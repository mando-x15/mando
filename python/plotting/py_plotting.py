#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri May  7 20:23:53 2021

https://www.earthdatascience.org/courses/use-data-open-source-python/use-time-series-data-in-python/date-time-types-in-pandas-python/resample-time-series-data-pandas-python/


https://benalexkeen.com/resampling-time-series-data-with-pandas/

range = pd.date_range('2015-01-01', '2015-12-31', freq='15min')
df = pd.DataFrame(index = range)


@author: spfeife
"""

import docx

import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime
from datetime import timedelta
import seaborn as sns
import numpy as np


mydoc = docx.Document()

mydoc.add_paragraph("This is first paragraph of a MS Word file.")

mydoc.add_heading("This is level 1 heading", 0)

# Handle date time conversions between pandas and matplotlib
from pandas.plotting import register_matplotlib_converters
register_matplotlib_converters()

# Use white grid plot background from seaborn
sns.set(font_scale=1.5, style="whitegrid")


file_path = '~/Desktop/Python/plotting/colorado-flood/precipitation/805325-precip-daily-2003-2013.csv'

# Import data using datetime and no data value
data_hourly = pd.read_csv(file_path,
                   parse_dates=['DATE'],
                   index_col=['DATE'],
                   na_values=['999.99'])

# Create figure and plot space
fig, ax = plt.subplots(figsize=(10, 6))

# Add x-axis and y-axis
ax.scatter(data_hourly.index.values,
           data_hourly['HPCP'],
           color='purple')

# Set title and labels for axes
ax.set(xlabel="Date",
       ylabel="Precipitation (inches)",
       title="Hourly Precipitation - Boulder Station\n 2003-2013")

#plt.show()
plt.savefig('plot1.png', dpi=300, bbox_inches='tight')
mydoc.add_picture("plot1.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))

#%%

# Resample to daily precip sum and save as new dataframe
data_daily = data_hourly.resample('D').sum()


# Create figure and plot space
fig, ax = plt.subplots(figsize=(10, 6))

# Add x-axis and y-axis
ax.scatter(data_daily.index.values,
           data_daily['HPCP'],
           color='purple')

# Set title and labels for axes
ax.set(xlabel="Date",
       ylabel="Precipitation (inches)",
       title="Daily Precipitation - Boulder Station\n 2003-2013")

#plt.show()
plt.savefig('plot2.png', dpi=300, bbox_inches='tight')
mydoc.add_picture("plot2.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))

#%%

# Resample to monthly precip sum and save as new dataframe
data_monthly = data_daily.resample('M').sum()

# Create figure and plot space
fig, ax = plt.subplots(figsize=(10, 6))

# Add x-axis and y-axis
ax.scatter(data_monthly.index.values, data_monthly['HPCP'], color='purple')

# Set title and labels for axes
ax.set(xlabel="Date",
       ylabel="Precipitation (inches)",
       title="Monthly Precipitation - Boulder Station\n 2003-2013")

plt.plot(data_monthly.index.values, data_monthly['HPCP'], color='darkred')

#plt.show()
plt.savefig('plot3.png', dpi=300, bbox_inches='tight')
mydoc.add_picture("plot3.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))

#%%

# Resample to yearly precip sum and save as new dataframe
data_yearly = data_hourly.resample('Y').sum()

# Create figure and plot space
fig, ax = plt.subplots(figsize=(10, 6))

# Add x-axis and y-axis
ax.scatter(data_yearly.index.values,
           data_yearly['HPCP'],
           color='purple')

plt.plot(data_yearly.index.values, data_yearly['HPCP'], color='darkred')

# Set title and labels for axes
ax.set(xlabel="Date",
       ylabel="Precipitation (inches)",
       title="Yearly Precipitation - Boulder Station\n 2003-2013")

#plt.show()

plt.savefig('plot4.png', dpi=300, bbox_inches='tight')
mydoc.add_picture("plot4.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))

mydoc.save("my_written_file.docx")

#%%

range = pd.date_range('2015-01-01', '2015-12-31', freq='D')
df = pd.DataFrame(index = range)
df['count'] = 0




#%%




#week_groups = data.groupby([data['date_year'], data['date_week']])['value'].count()
#week_groups.plot(kind='bar', figsize=(10, 5), legend=None)

#weekly_data = data.groupby("Channel").resample('W-Wed', label='right', closed = 'right', on='Day').sum().reset_index().sort_values(by='Day')

# set the index to be the date for the data
#data1 = data.sort_values('date').set_index('date').resample('W').value.count()

# create bar chart and update the date format for the weeks
#ax = data1.plot(kind='line',figsize=(10, 5), legend=None)
#ax.set_xticklabels(data1.index.strftime('%Y-%m-%d'), rotation=90)
#plt.xlabel('Date - Week Starting')
